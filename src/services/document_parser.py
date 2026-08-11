import io
from pathlib import Path
from typing import Dict, Any

def extract_text_from_file(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Extract clean text and metadata from uploaded PDF, Word, or text files."""
    ext = Path(filename).suffix.lower()
    text = ""
    pages_or_items = 0
    file_type = "text"

    try:
        if ext == ".pdf":
            file_type = "pdf"
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            pages_or_items = len(reader.pages)
            page_texts = []
            for idx, page in enumerate(reader.pages, 1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    page_texts.append(f"--- [Page {idx}] ---\n{page_text.strip()}")
            text = "\n\n".join(page_texts)

        elif ext in (".docx", ".doc"):
            file_type = "word"
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables if present
            table_texts = []
            for t_idx, table in enumerate(doc.tables, 1):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    table_texts.append(f"[Table {t_idx}]\n" + "\n".join(rows))

            all_parts = paragraphs + table_texts
            pages_or_items = len(all_parts)
            text = "\n\n".join(all_parts)

        elif ext in (".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"):
            file_type = "image"
            import base64
            b64_data = base64.b64encode(file_bytes).decode("utf-8")
            mime_sub = "jpeg" if ext in (".jpg", ".jpeg") else ext.lstrip(".")
            image_url = f"data:image/{mime_sub};base64,{b64_data}"
            text = f"[Image attached: {filename}]"
            pages_or_items = 1
            return {
                "success": True,
                "filename": filename,
                "text": text,
                "file_type": "image",
                "image_url": image_url,
                "items_count": 1,
                "size": len(file_bytes),
                "is_truncated": False
            }

        else:
            # Plain text, Markdown, CSV, JSON, Python, etc.
            file_type = "text"
            text = file_bytes.decode("utf-8", errors="replace")
            pages_or_items = len(text.splitlines())


    except Exception as e:
        return {
            "success": False,
            "filename": filename,
            "error": f"Failed to parse document: {str(e)}",
            "text": "",
            "file_type": file_type,
            "size": len(file_bytes)
        }

    # Limit text size if exceptionally huge (e.g. cap at 50,000 characters)
    is_truncated = False
    if len(text) > 60000:
        text = text[:60000] + "\n\n... [Content truncated to first 60,000 characters] ..."
        is_truncated = True

    return {
        "success": True,
        "filename": filename,
        "text": text.strip(),
        "file_type": file_type,
        "items_count": pages_or_items,
        "size": len(file_bytes),
        "is_truncated": is_truncated
    }
